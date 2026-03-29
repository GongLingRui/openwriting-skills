#!/usr/bin/env python3
"""
Markdown报告 → 公文风格DOCX转换器

将AI生成的Markdown格式报告转换为符合央企公文规范的Word文档。

用法：
  # 从文件转换
  python3 generate_docx.py --input report.md --output report.docx

  # 从stdin转换（管道模式，可与LLM输出串联）
  echo "报告内容" | python3 generate_docx.py --output report.docx

  # 指定标题和日期
  python3 generate_docx.py --input report.md --output report.docx --title "周进度简报" --date "2026年3月12日"
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ====== Emoji → 公文友好符号映射 ======
EMOJI_MAP = {
    '🔴': ('●', RGBColor(200, 30, 30)),     # 红色圆点
    '🟡': ('●', RGBColor(200, 160, 0)),     # 黄色圆点
    '🟢': ('●', RGBColor(30, 160, 30)),     # 绿色圆点
    '⚠️': ('▲', RGBColor(200, 130, 0)),    # 橙色三角
    '❌': ('✕', RGBColor(200, 30, 30)),     # 红色叉
    '✅': ('✓', RGBColor(30, 160, 30)),     # 绿色勾
    '📋': ('', None),                       # 移除装饰性emoji
    '📊': ('', None),
    '📄': ('', None),
}

# 构建emoji检测正则
_EMOJI_PATTERN = re.compile('(' + '|'.join(re.escape(e) for e in EMOJI_MAP.keys()) + ')')


# ====== 公文样式配置 ======
STYLE_CONFIG = {
    'page': {
        'top_margin': Cm(3.7),     # 上边距
        'bottom_margin': Cm(3.5),  # 下边距
        'left_margin': Cm(2.8),    # 左边距
        'right_margin': Cm(2.6),   # 右边距
    },
    'title': {
        'font_name': '方正小标宋简体',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(22),
        'color': RGBColor(0, 0, 0),
        'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'space_before': Pt(0),
        'space_after': Pt(20),
        'line_spacing': Pt(28),
    },
    'subtitle': {
        'font_name': '楷体_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(16),
        'color': RGBColor(0, 0, 0),
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'space_before': Pt(0),
        'space_after': Pt(15),
    },
    'h1': {
        'font_name': '黑体',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(16),
        'bold': True,
        'color': RGBColor(0, 0, 0),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'space_before': Pt(15),
        'space_after': Pt(8),
        'line_spacing': Pt(28),
    },
    'h2': {
        'font_name': '楷体_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(15),
        'bold': True,
        'color': RGBColor(0, 0, 0),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'space_before': Pt(10),
        'space_after': Pt(6),
        'line_spacing': Pt(28),
    },
    'h3': {
        'font_name': '仿宋_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(14),
        'bold': True,
        'color': RGBColor(0, 0, 0),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'space_before': Pt(8),
        'space_after': Pt(4),
        'line_spacing': Pt(28),
    },
    'body': {
        'font_name': '仿宋_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(14),
        'color': RGBColor(0, 0, 0),
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'first_line_indent': Cm(0.74),  # 首行缩进2字符
        'space_before': Pt(0),
        'space_after': Pt(3),
        'line_spacing': Pt(28),
    },
    'table': {
        'font_name': '仿宋_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(10.5),
        'header_font_name': '黑体',
        'header_font_size': Pt(10.5),
        'header_bold': True,
        'header_bg': RGBColor(218, 227, 243),  # 浅蓝背景
    },
    'footer': {
        'font_name': '仿宋_GB2312',
        'font_name_ascii': 'Times New Roman',
        'font_size': Pt(9),
        'color': RGBColor(128, 128, 128),
    },
}


def apply_font(run, style_key):
    """应用字体样式到 run"""
    style = STYLE_CONFIG[style_key]
    run.font.size = style.get('font_size', Pt(14))
    run.font.bold = style.get('bold', False)
    run.font.color.rgb = style.get('color', RGBColor(0, 0, 0))
    # 设置中文字体
    font_name = style.get('font_name', '仿宋_GB2312')
    run.font.name = style.get('font_name_ascii', 'Times New Roman')
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def apply_paragraph_format(paragraph, style_key):
    """应用段落格式"""
    style = STYLE_CONFIG[style_key]
    pf = paragraph.paragraph_format
    pf.alignment = style.get('alignment', WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf.space_before = style.get('space_before', Pt(0))
    pf.space_after = style.get('space_after', Pt(3))
    if 'line_spacing' in style:
        pf.line_spacing = style['line_spacing']
    if 'first_line_indent' in style:
        pf.first_line_indent = style['first_line_indent']


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_text_with_emoji(paragraph, text, style_key, font_name_override=None):
    """
    向段落添加文本，自动将emoji替换为带颜色的符号。
    文本会按emoji边界拆分为多个run，emoji部分使用对应颜色。
    """
    parts = _EMOJI_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part in EMOJI_MAP:
            symbol, color = EMOJI_MAP[part]
            if not symbol:  # 移除装饰性emoji
                continue
            run = paragraph.add_run(symbol)
            apply_font(run, style_key)
            if font_name_override:
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), font_name_override)
            if color:
                run.font.color.rgb = color
        else:
            run = paragraph.add_run(part)
            apply_font(run, style_key)
            if font_name_override:
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), font_name_override)


def clean_text(text):
    """清理Markdown标记符号"""
    # 去除粗体标记
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 去除斜体标记
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 去除行内代码
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 去除emoji（保留状态灯）
    # 去除链接
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


def parse_markdown(content):
    """
    解析Markdown内容为结构化元素列表。
    返回: [(type, data), ...]
    type: 'title', 'h1', 'h2', 'h3', 'paragraph', 'table', 'separator', 'list_item'
    """
    elements = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行
        if not line.strip():
            i += 1
            continue

        # 水平分隔线
        if re.match(r'^-{3,}$|^\*{3,}$|^={3,}$', line.strip()):
            elements.append(('separator', None))
            i += 1
            continue

        # 标题 #
        m = re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            title_text = clean_text(m.group(2))
            if level == 1:
                elements.append(('h1', title_text))
            elif level == 2:
                elements.append(('h2', title_text))
            elif level == 3:
                elements.append(('h3', title_text))
            i += 1
            continue

        # 表格检测
        if '|' in line and i + 1 < len(lines):
            # 检查是否是表格（至少有表头和分隔行）
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j].strip():
                table_lines.append(lines[j].strip())
                j += 1

            if len(table_lines) >= 2:
                # 解析表格
                rows = []
                for tl in table_lines:
                    # 跳过分隔行 (|---|---|)
                    if re.match(r'^\|[\s\-:|]+\|?$', tl):
                        continue
                    cells = [clean_text(c.strip()) for c in tl.split('|')]
                    # 去除首尾空元素
                    cells = [c for c in cells if c or cells.index(c) not in (0, len(cells) - 1)]
                    if not cells:
                        continue
                    # 去掉首尾可能的空字符串
                    while cells and not cells[0]:
                        cells.pop(0)
                    while cells and not cells[-1]:
                        cells.pop()
                    if cells:
                        rows.append(cells)
                if rows:
                    elements.append(('table', rows))
                i = j
                continue

        # 列表项
        m = re.match(r'^(\s*)([\-\*]|\d+\.)\s+(.+)$', line)
        if m:
            indent = len(m.group(1))
            marker = m.group(2)
            text = clean_text(m.group(3))
            level = indent // 2  # 0, 1, 2...
            is_ordered = bool(re.match(r'\d+\.', marker))
            elements.append(('list_item', {'text': text, 'level': level, 'ordered': is_ordered}))
            i += 1
            continue

        # 引用块
        if line.strip().startswith('>'):
            text = clean_text(line.strip().lstrip('> '))
            elements.append(('quote', text))
            i += 1
            continue

        # 普通段落
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].rstrip()
            if not next_line.strip():
                break
            if re.match(r'^#{1,3}\s', next_line):
                break
            if '|' in next_line and j + 1 < len(lines) and '|' in lines[j + 1]:
                break
            if re.match(r'^(\s*)([\-\*]|\d+\.)\s', next_line):
                break
            if re.match(r'^-{3,}$|^\*{3,}$|^={3,}$', next_line.strip()):
                break
            para_lines.append(next_line)
            j += 1

        full_text = clean_text(' '.join(para_lines))
        if full_text:
            elements.append(('paragraph', full_text))
        i = j

    return elements


def create_docx(elements, title=None, subtitle=None, date_str=None, footer_text=None):
    """
    从结构化元素创建DOCX文档。
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    page = STYLE_CONFIG['page']
    section.top_margin = page['top_margin']
    section.bottom_margin = page['bottom_margin']
    section.left_margin = page['left_margin']
    section.right_margin = page['right_margin']

    # 文档标题
    if title:
        p = doc.add_paragraph()
        apply_paragraph_format(p, 'title')
        run = p.add_run(title)
        apply_font(run, 'title')

    # 副标题（日期等）
    if subtitle:
        p = doc.add_paragraph()
        apply_paragraph_format(p, 'subtitle')
        run = p.add_run(subtitle)
        apply_font(run, 'subtitle')

    if date_str:
        p = doc.add_paragraph()
        apply_paragraph_format(p, 'subtitle')
        run = p.add_run(f"（{date_str}）")
        apply_font(run, 'subtitle')

    # 渲染元素
    for elem_type, data in elements:
        if elem_type == 'h1':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'h1')
            add_text_with_emoji(p, data, 'h1')

        elif elem_type == 'h2':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'h2')
            add_text_with_emoji(p, data, 'h2')

        elif elem_type == 'h3':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'h3')
            add_text_with_emoji(p, data, 'h3')

        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'body')
            add_text_with_emoji(p, data, 'body')

        elif elem_type == 'quote':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'body')
            p.paragraph_format.left_indent = Cm(1.0)
            add_text_with_emoji(p, data, 'body')
            # 非emoji的run设为灰色
            for run in p.runs:
                if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0, 0, 0):
                    run.font.color.rgb = RGBColor(80, 80, 80)

        elif elem_type == 'list_item':
            p = doc.add_paragraph()
            apply_paragraph_format(p, 'body')
            p.paragraph_format.first_line_indent = None
            indent = Cm(0.74 + data['level'] * 0.74)
            p.paragraph_format.left_indent = indent
            text = data['text']
            add_text_with_emoji(p, text, 'body')

        elif elem_type == 'table':
            rows = data
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            # 统一列数
            for r in rows:
                while len(r) < num_cols:
                    r.append('')

            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 设置表格边框
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                borders.append(border)
            tblPr.append(borders)

            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ts = STYLE_CONFIG['table']

                    if row_idx == 0:
                        # 表头样式 - 无emoji
                        run = p.add_run(str(cell_text))
                        run.font.size = ts['header_font_size']
                        run.font.bold = ts['header_bold']
                        run.font.name = ts.get('font_name_ascii', 'Times New Roman')
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.insert(0, rFonts)
                        rFonts.set(qn('w:eastAsia'), ts['header_font_name'])
                        set_cell_shading(cell, ts['header_bg'])
                    else:
                        # 数据行 - 支持emoji替换为彩色符号
                        add_text_with_emoji(
                            p, str(cell_text), 'table',
                            font_name_override=ts['font_name']
                        )

            # 表格后加空段落
            doc.add_paragraph()

        elif elem_type == 'separator':
            # 不插入分隔线，公文不用
            pass

    # 页脚
    if footer_text:
        footer_section = doc.sections[0].footer
        p = footer_section.paragraphs[0] if footer_section.paragraphs else footer_section.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(footer_text)
        apply_font(run, 'footer')

    return doc


def extract_title_from_content(content):
    """从Markdown内容中提取标题"""
    for line in content.split('\n'):
        line = line.strip()
        m = re.match(r'^#\s+(.+)$', line)
        if m:
            return clean_text(m.group(1))
    return None


def main():
    parser = argparse.ArgumentParser(description='将Markdown报告转换为公文风格DOCX文档')
    parser.add_argument('--input', '-i', type=str, help='输入Markdown文件路径（不指定则从stdin读取）')
    parser.add_argument('--output', '-o', type=str, required=True, help='输出DOCX文件路径')
    parser.add_argument('--title', '-t', type=str, help='文档标题（不指定则从内容中提取）')
    parser.add_argument('--date', '-d', type=str, help='文档日期')
    parser.add_argument('--footer', type=str, default='以上分析基于填报数据，管理决策请结合实际情况',
                        help='页脚文字')

    args = parser.parse_args()

    # 读取内容
    if args.input:
        if not os.path.exists(args.input):
            print(f"Error: File not found: {args.input}", file=sys.stderr)
            sys.exit(1)
        with open(args.input, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        content = sys.stdin.read()

    if not content.strip():
        print("Error: Empty content", file=sys.stderr)
        sys.exit(1)

    # 清理 <think>...</think> 块
    content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)

    # 提取或使用指定标题
    title = args.title or extract_title_from_content(content)

    # 如果从内容中提取了标题，从正文中移除（避免重复）
    if title and not args.title:
        content = re.sub(r'^#\s+.+\n', '', content, count=1)

    # 解析
    elements = parse_markdown(content)

    # 生成DOCX
    doc = create_docx(
        elements,
        title=title,
        date_str=args.date,
        footer_text=args.footer
    )

    # 保存
    doc.save(args.output)
    print(f"✅ DOCX文档已生成: {args.output}")
    print(f"   标题: {title or '(无)'}")
    print(f"   页数: 约{len(elements) // 8 + 1}页（估算）")


if __name__ == '__main__':
    main()
